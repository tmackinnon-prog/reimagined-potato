"""TPGP Generator - Flask backend.

Serves the wizard front end and a single /api/generate endpoint that:
  1. parses any uploaded files (PDF / DOCX / DOC / TXT / MD),
  2. fetches and extracts text from any supplied URLs,
  3. builds a prompt from the teacher's TQS self-reflection + that context,
  4. asks Claude to draft a TPGP as structured JSON (via forced tool use),
  5. fills the Third Schools Word template with the result, and
  6. returns the plan JSON (for the on-page preview) plus the .docx as
     base64 (for the browser to download).

Nothing is persisted server-side; uploaded files and the API key exist
only for the duration of a single request.
"""
import base64
import io
import json
import os
import re
import subprocess
import tempfile

from flask import Flask, jsonify, request, send_from_directory

import requests
from bs4 import BeautifulSoup
import pdfplumber
from docx import Document as DocxDocument

import docx_builder
from indicators import INDICATORS, RATING_LABELS
from tpgp_reference import COMPETENCIES, TPGP_GUIDANCE

APP_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__, static_folder=None)
app.config["MAX_CONTENT_LENGTH"] = 30 * 1024 * 1024  # 30 MB of uploads per request

ANTHROPIC_API_URL = "https://api.anthropic.com/v1/messages"
ANTHROPIC_VERSION = "2023-06-01"

MAX_FILE_CHARS = 12000
MAX_URL_CHARS = 6000
ALLOWED_FILE_EXT = {".pdf", ".docx", ".doc", ".txt", ".md"}


# --------------------------------------------------------------- front end

@app.route("/")
def index():
    return send_from_directory(APP_DIR, "index.html")


# ------------------------------------------------------------- extraction

def _truncate(s, n):
    s = (s or "").strip()
    return s if len(s) <= n else s[:n] + "\n…[truncated]"


def extract_pdf_text(raw_bytes):
    text = []
    with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
        for page in pdf.pages[:40]:
            t = page.extract_text() or ""
            if t:
                text.append(t)
    return "\n".join(text)


def extract_docx_text(raw_bytes):
    d = DocxDocument(io.BytesIO(raw_bytes))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = " | ".join(c.text for c in row.cells)
            if cells.strip():
                parts.append(cells)
    return "\n".join(parts)


def extract_doc_text(raw_bytes, filename):
    """Legacy .doc: try antiword first, fall back to a LibreOffice conversion."""
    with tempfile.TemporaryDirectory() as td:
        src = os.path.join(td, filename)
        with open(src, "wb") as f:
            f.write(raw_bytes)
        try:
            result = subprocess.run(["antiword", src], capture_output=True, timeout=30)
            if result.returncode == 0 and result.stdout:
                return result.stdout.decode("utf-8", "ignore")
        except (FileNotFoundError, subprocess.SubprocessError):
            pass
        try:
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "docx", "--outdir", td, src],
                capture_output=True, timeout=60,
            )
            converted = os.path.join(td, os.path.splitext(filename)[0] + ".docx")
            if os.path.exists(converted):
                with open(converted, "rb") as f:
                    return extract_docx_text(f.read())
        except (FileNotFoundError, subprocess.SubprocessError):
            pass
    return ""


def extract_text_file(raw_bytes):
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return ""


def parse_uploaded_files(files):
    parts = []
    for f in files:
        name = f.filename or "file"
        ext = os.path.splitext(name)[1].lower()
        if ext not in ALLOWED_FILE_EXT:
            continue
        raw = f.read()
        if not raw:
            continue
        try:
            if ext == ".pdf":
                text = extract_pdf_text(raw)
            elif ext == ".docx":
                text = extract_docx_text(raw)
            elif ext == ".doc":
                text = extract_doc_text(raw, name)
            else:
                text = extract_text_file(raw)
        except Exception:
            text = ""
        text = _truncate(text, MAX_FILE_CHARS)
        if text:
            parts.append(f"### Uploaded file: {name}\n{text}")
    return parts


def fetch_url_text(url):
    try:
        resp = requests.get(
            url, timeout=12,
            headers={"User-Agent": "Mozilla/5.0 (compatible; TPGP-Generator/1.0)"},
        )
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "noscript", "svg"]):
            tag.decompose()
        text = soup.get_text("\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        return _truncate(text, MAX_URL_CHARS)
    except Exception:
        return None


def parse_urls(urls):
    parts, failed = [], []
    for u in urls:
        u = (u or "").strip()
        if not u:
            continue
        text = fetch_url_text(u)
        if text:
            parts.append(f"### Web page: {u}\n{text}")
        else:
            failed.append(u)
    return parts, failed


# ------------------------------------------------------------- prompt build

def build_reflection_summary(reflection):
    reflection = reflection or {}
    ti = reflection.get("teacherInfo") or {}
    ratings = reflection.get("ratings") or {}
    comments = reflection.get("comments") or {}
    goals = reflection.get("goals") or []
    written = reflection.get("reflections") or {}

    lines = []
    lines.append("TEACHER DETAILS")
    lines.append(f"- Name: {ti.get('name') or '(not given)'}")
    lines.append(f"- School / Program: {ti.get('school') or '(not given)'}")
    lines.append(f"- Subjects / Grades (Position): {ti.get('subjects') or '(not given)'}")
    lines.append(f"- Years of experience: {ti.get('experience') or '(not given)'}")
    lines.append("")
    lines.append("TQS SELF-REFLECTION RATINGS (1=Rarely/Not Yet .. 5=Exemplary)")
    for comp_id in (1, 2, 3, 4, 5, 6):
        comp = COMPETENCIES[comp_id]
        comp_ratings = ratings.get(str(comp_id)) or ratings.get(comp_id) or []
        rated_lines = []
        for i, (code, text) in enumerate(INDICATORS[comp_id]):
            val = comp_ratings[i] if i < len(comp_ratings) else 0
            note = (comments.get(f"{comp_id}-{i}") or "").strip()
            if val or note:
                label = RATING_LABELS.get(val, "not rated")
                extra = f" | note: {note}" if note else ""
                rated_lines.append(f"    {code} ({label}={val or '-'}): {text}{extra}")
        if rated_lines:
            lines.append(f"  Competency {comp_id} - {comp['title']}:")
            lines.extend(rated_lines)
    lines.append("")

    if goals:
        lines.append("TEACHER'S OWN DRAFT GOALS (build on these, refine as needed)")
        for g in goals:
            lines.append(
                f"  - Competency: {g.get('competency','')} | Change: {g.get('action','')} | "
                f"By: {g.get('when','')} | Support needed: {g.get('support','')}"
            )
        lines.append("")

    if any((written or {}).values()):
        lines.append("WRITTEN REFLECTIONS")
        if written.get("strengths"):
            lines.append(f"  - Things I do well: {written['strengths']}")
        if written.get("strengthen"):
            lines.append(f"  - TQS areas I want to strengthen: {written['strengthen']}")
        if written.get("newStrategy"):
            lines.append(f"  - A new strategy/approach I want to try: {written['newStrategy']}")
        if written.get("support"):
            lines.append(f"  - Professional learning/support I'd find helpful: {written['support']}")
        if written.get("additional"):
            lines.append(f"  - Additional thoughts: {written['additional']}")
        lines.append("")

    return "\n".join(lines)


def build_user_prompt(meta, file_parts, url_parts, failed_urls):
    num_goals = int(meta.get("numGoals") or 3)
    reflection_summary = build_reflection_summary(meta.get("reflection"))
    notes = (meta.get("notes") or "").strip()
    school_year = (meta.get("schoolYear") or "").strip()

    parts = [reflection_summary]
    if notes:
        parts.append(f"ADDITIONAL CONTEXT FROM THE TEACHER\n{notes}\n")
    if file_parts:
        parts.append("SUPPORTING FILES\n" + "\n\n".join(file_parts) + "\n")
    if url_parts:
        parts.append("REFERENCED WEB PAGES\n" + "\n\n".join(url_parts) + "\n")
    if failed_urls:
        parts.append(
            "(Note: the following links could not be read and were skipped: "
            + ", ".join(failed_urls) + ")\n"
        )

    parts.append(
        f"TASK\nDraft a Teacher Professional Growth Plan with exactly {num_goals} goal(s) "
        f"for the {school_year or 'upcoming'} school year, using the submit_tpgp tool. "
        "Ground it in the reflection above and follow the drafting guidance in the system prompt."
    )
    return "\n\n".join(parts)


PLAN_TOOL = {
    "name": "submit_tpgp",
    "description": "Submit the completed draft Teacher Professional Growth Plan.",
    "input_schema": {
        "type": "object",
        "properties": {
            "plan_month_year": {
                "type": "string",
                "description": "e.g. 'September 2026' - the plan's start month/year.",
            },
            "rationale": {
                "type": "string",
                "description": "2-4 sentences to the teacher explaining how their reflection shaped these goals.",
            },
            "goals": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "title": {"type": "string"},
                        "competencies_of_focus": {"type": "string"},
                        "timeframe": {"type": "string", "description": "e.g. 'June, 2027'"},
                        "short_term_objectives": {"type": "array", "items": {"type": "string"}},
                        "strategies": {"type": "array", "items": {"type": "string"}},
                        "measures": {"type": "array", "items": {"type": "string"}},
                        "resources": {"type": "array", "items": {"type": "string"}},
                    },
                    "required": [
                        "title", "competencies_of_focus", "timeframe",
                        "short_term_objectives", "strategies", "measures", "resources",
                    ],
                },
            },
        },
        "required": ["plan_month_year", "rationale", "goals"],
    },
}


def call_claude(meta, user_prompt):
    api_key = (meta.get("apiKey") or "").strip() or os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        raise GenerationError(
            "No Anthropic API key found. Paste one into the API Key field on Step 4, "
            "or set the ANTHROPIC_API_KEY environment variable on the server."
        )
    model = (meta.get("model") or "claude-sonnet-5").strip()

    resp = requests.post(
        ANTHROPIC_API_URL,
        headers={
            "x-api-key": api_key,
            "anthropic-version": ANTHROPIC_VERSION,
            "content-type": "application/json",
        },
        json={
            "model": model,
            "max_tokens": 4096,
            "system": TPGP_GUIDANCE,
            "tools": [PLAN_TOOL],
            "tool_choice": {"type": "tool", "name": "submit_tpgp"},
            "messages": [{"role": "user", "content": user_prompt}],
        },
        timeout=90,
    )
    if resp.status_code != 200:
        try:
            detail = resp.json().get("error", {}).get("message", resp.text)
        except Exception:
            detail = resp.text
        raise GenerationError(f"Claude API error ({resp.status_code}): {detail}")

    data = resp.json()
    for block in data.get("content", []):
        if block.get("type") == "tool_use" and block.get("name") == "submit_tpgp":
            return block.get("input") or {}
    raise GenerationError("Claude did not return a structured plan. Please try again.")


class GenerationError(Exception):
    pass


def sanitize_filename(s):
    s = re.sub(r"[^A-Za-z0-9 _.\-]", "", s or "").strip()
    return re.sub(r"\s+", " ", s)


@app.route("/api/generate", methods=["POST"])
def api_generate():
    try:
        meta_raw = request.form.get("meta", "{}")
        try:
            meta = json.loads(meta_raw)
        except json.JSONDecodeError:
            return jsonify({"error": "Malformed request."}), 400

        files = request.files.getlist("files")
        file_parts = parse_uploaded_files(files)

        urls = meta.get("urls") or []
        url_parts, failed_urls = parse_urls(urls)

        user_prompt = build_user_prompt(meta, file_parts, url_parts, failed_urls)
        plan = call_claude(meta, user_prompt)

        num_goals = int(meta.get("numGoals") or 3)
        plan["goals"] = (plan.get("goals") or [])[:num_goals]

        docx_bytes = docx_builder.build_docx(plan, meta)

        teacher_name = ((meta.get("reflection") or {}).get("teacherInfo") or {}).get("name") or "Teacher"
        school_year = meta.get("schoolYear") or ""
        filename = sanitize_filename(f"TPGP - {teacher_name} - {school_year}") + ".docx"

        return jsonify({
            "plan": plan,
            "docx_base64": base64.b64encode(docx_bytes).decode("ascii"),
            "filename": filename,
        })
    except GenerationError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": f"Unexpected server error: {e}"}), 500


if __name__ == "__main__":
    port = int(os.environ.get("TPGP_PORT", "8200"))
    app.run(host="0.0.0.0", port=port)
