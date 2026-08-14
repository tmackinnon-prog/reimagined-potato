"""Fills the Third Schools TPGP Word template with a generated plan.

Falls back to building a plain document from scratch if the template file
is missing or fails to open, so the app still produces a usable .docx.
"""
import os
import re
from io import BytesIO

from docx import Document
from docx.shared import Pt

from tpgp_reference import COMPETENCIES

TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tpgp_template.docx")

_GOAL_RE = re.compile(r"^Goal\s+(\d+)$")
_NUM_RE = re.compile(r"^(\d+)\.\s*(.*)$")


def build_docx(plan, meta):
    """Returns raw .docx bytes for the given plan dict and request meta dict."""
    if os.path.exists(TEMPLATE_PATH):
        try:
            return _build_from_template(plan, meta)
        except Exception:
            pass
    return _build_from_scratch(plan, meta)


# ---------------------------------------------------------------- template

def _build_from_template(plan, meta):
    doc = Document(TEMPLATE_PATH)
    teacher = (meta.get("reflection") or {}).get("teacherInfo") or {}
    goals = plan.get("goals") or []
    school_year = (meta.get("schoolYear") or "").strip()
    start_year, end_year = _parse_school_year(school_year)

    _set_title_subtitle(doc, plan.get("plan_month_year") or _default_month_year(start_year))
    _fill_label_line(doc, "Name:", teacher.get("name") or "")
    _fill_label_line(doc, "Position:", teacher.get("subjects") or "")
    _update_review_dates(doc, start_year, end_year)

    for n in (1, 2, 3):
        start, end = _goal_bounds(doc, n)
        if start is None:
            continue
        goal = goals[n - 1] if n - 1 < len(goals) else None
        if goal is None:
            _remove_range(doc, start, end)
            continue
        _fill_goal_block(doc, n, goal, end_year)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _parse_school_year(school_year):
    m = re.match(r"^\s*(\d{4})\s*-\s*(\d{2,4})\s*$", school_year or "")
    if not m:
        return None, None
    start = int(m.group(1))
    end_raw = m.group(2)
    end = int(end_raw) if len(end_raw) == 4 else (start // 100) * 100 + int(end_raw)
    return start, end


def _default_month_year(start_year):
    if start_year:
        return f"September {start_year}"
    return "September"


def _set_title_subtitle(doc, month_year):
    for p in doc.paragraphs:
        if p.text.strip() and re.match(r"^(January|February|March|April|May|June|July|August|"
                                        r"September|October|November|December)\s+\d{4}$", p.text.strip()):
            _set_text_keep_first_run(p, month_year)
            return


def _fill_label_line(doc, label, value):
    if not value:
        return
    for p in doc.paragraphs:
        if p.text.strip().startswith(label):
            _append_run(p, " " + value.strip())
            return


def _update_review_dates(doc, start_year, end_year):
    if not start_year or not end_year:
        return
    replacements = {
        r"^November\s+\d{4}:$": f"November {start_year}:",
        r"^Jan\s+\d{4}:$": f"Jan {end_year}:",
        r"^June\s+\d{4}:$": f"June {end_year}:",
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        for pattern, replacement in replacements.items():
            if re.match(pattern, t):
                _set_text_keep_first_run(p, replacement)


def _goal_bounds(doc, n):
    paras = doc.paragraphs
    start = None
    for i, p in enumerate(paras):
        if p.text.strip() == f"Goal {n}":
            start = i
            break
    if start is None:
        return None, None
    end = len(paras)
    for j in range(start + 1, len(paras)):
        t = paras[j].text.strip()
        if _GOAL_RE.match(t) and t != f"Goal {n}":
            end = j
            break
        if t == "Reflection on Success:":
            end = j
            break
    return start, end


def _fill_goal_block(doc, n, goal, end_year):
    # Fill everything that relies on locating this block by its exact
    # "Goal N" heading text *before* renaming that heading to "Goal N: Title" -
    # renaming it first would break every subsequent text-based lookup.
    _fill_label_line_within(doc, n, "Competencies of Focus:", goal.get("competencies_of_focus") or "")
    _fill_timeframe(doc, n, goal.get("timeframe") or (f"June, {end_year}" if end_year else ""))

    last_strategy_para = _fill_numbered_section(doc, n, "Short Term Objectives:", goal.get("short_term_objectives") or [])
    last_strategy_para = _fill_numbered_section(doc, n, "Strategies:", goal.get("strategies") or []) or last_strategy_para

    _insert_extra_sections(doc, n, [
        ("Measures / Indicators of Success:", goal.get("measures") or []),
        ("Resources:", goal.get("resources") or []),
    ])

    title = (goal.get("title") or "").strip()
    for p in doc.paragraphs:
        if p.text.strip() == f"Goal {n}":
            _set_text_keep_first_run(p, f"Goal {n}: {title}" if title else f"Goal {n}")
            break


def _fill_label_line_within(doc, n, label, value):
    if not value:
        return
    start, end = _goal_bounds(doc, n)
    if start is None:
        return
    for p in doc.paragraphs[start:end]:
        if p.text.strip().rstrip(":") == label.rstrip(":"):
            _append_run(p, " " + value.strip())
            return


def _fill_timeframe(doc, n, value):
    if not value:
        return
    start, end = _goal_bounds(doc, n)
    if start is None:
        return
    for p in doc.paragraphs[start:end]:
        if p.text.strip().startswith("By "):
            text = value.strip()
            if not text.lower().startswith("by "):
                text = "By " + text
            _set_text_keep_first_run(p, text)
            return


def _fill_numbered_section(doc, n, heading_text, items):
    """Fills the '1. ' / '2. ' lines that already follow a heading paragraph
    inside goal n's block, adding extra numbered lines if there are more
    than two items. Returns the last paragraph written, or None."""
    start, end = _goal_bounds(doc, n)
    if start is None:
        return None
    paras = doc.paragraphs
    heading_idx = None
    for i in range(start, end):
        if paras[i].text.strip() == heading_text:
            heading_idx = i
            break
    if heading_idx is None:
        return None

    slot_idxs = []
    for i in range(heading_idx + 1, end):
        if _NUM_RE.match(paras[i].text.strip()):
            slot_idxs.append(i)
        elif slot_idxs and paras[i].text.strip():
            break
    if not slot_idxs:
        return None

    items = [str(x).strip() for x in items if str(x).strip()]
    last_para = None
    for i, idx in enumerate(slot_idxs):
        p = doc.paragraphs[idx]
        if i < len(items):
            _set_text_keep_first_run(p, f"{i + 1}. {items[i]}")
        else:
            _set_text_keep_first_run(p, f"{i + 1}. ")
        last_para = p

    if len(items) > len(slot_idxs):
        anchor = doc.paragraphs[slot_idxs[-1]]
        for i in range(len(slot_idxs), len(items)):
            new_p = anchor.insert_paragraph_before(f"{i + 1}. {items[i]}")
            _match_run_format(new_p, anchor)
            anchor._p.addnext(new_p._p)
            anchor = new_p
            last_para = new_p

    return last_para


def _insert_extra_sections(doc, n, sections):
    start, end = _goal_bounds(doc, n)
    if start is None:
        return
    boundary = doc.paragraphs[end] if end < len(doc.paragraphs) else None
    if boundary is None:
        return

    strategies_heading = None
    for p in doc.paragraphs[start:end]:
        if p.text.strip() == "Strategies:":
            strategies_heading = p
            break

    for heading_text, items in sections:
        items = [str(x).strip() for x in items if str(x).strip()]
        if not items:
            continue
        heading_p = boundary.insert_paragraph_before(heading_text)
        if strategies_heading is not None:
            heading_p.style = strategies_heading.style
            src_run = strategies_heading.runs[0] if strategies_heading.runs else None
            for r in heading_p.runs:
                if src_run is not None:
                    r.bold = src_run.bold
                    r.font.name = src_run.font.name
                    if src_run.font.size:
                        r.font.size = src_run.font.size
        for i, item in enumerate(items):
            item_p = boundary.insert_paragraph_before(f"{i + 1}. {item}")
            item_p.style = doc.styles["Normal"]
            for run in item_p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(14)


def _remove_range(doc, start, end):
    paras = doc.paragraphs[start:end]
    for p in paras:
        el = p._p
        el.getparent().remove(el)


def _set_text_keep_first_run(p, text):
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""


def _append_run(p, text):
    if p.runs:
        base = p.runs[-1]
        r = p.add_run(text)
        r.font.name = base.font.name
        if base.font.size:
            r.font.size = base.font.size
        r.bold = base.bold
    else:
        p.add_run(text)


def _match_run_format(target_p, source_p):
    if not source_p.runs or not target_p.runs:
        return
    src = source_p.runs[0]
    for r in target_p.runs:
        r.font.name = src.font.name
        if src.font.size:
            r.font.size = src.font.size
        r.bold = src.bold


# ------------------------------------------------------------- from scratch

def _build_from_scratch(plan, meta):
    doc = Document()
    teacher = (meta.get("reflection") or {}).get("teacherInfo") or {}

    doc.add_heading("Teacher Professional Growth Plan", level=1)
    doc.add_paragraph(plan.get("plan_month_year") or "")
    doc.add_paragraph(
        "Quality teaching occurs when the teacher's ongoing analysis of the context, and the "
        "teacher's decisions about which pedagogical knowledge and abilities to apply, result "
        "in optimum learning for all students."
    )
    doc.add_paragraph(f"Name: {teacher.get('name') or ''}")
    doc.add_paragraph(f"Position: {teacher.get('subjects') or ''}")

    doc.add_paragraph(
        "Your Teacher Professional Growth Plan should be developed with the Teaching Quality "
        "Standard's 6 Competencies in mind."
    )
    for i, c in COMPETENCIES.items():
        doc.add_paragraph(f"{i}. {c['title']}", style="List Bullet")

    if plan.get("rationale"):
        doc.add_heading("Rationale", level=2)
        doc.add_paragraph(plan["rationale"])

    for i, goal in enumerate(plan.get("goals") or [], start=1):
        doc.add_heading(f"Goal {i}: {goal.get('title', '')}", level=2)
        if goal.get("competencies_of_focus"):
            doc.add_paragraph(f"Competencies of Focus: {goal['competencies_of_focus']}")
        if goal.get("timeframe"):
            doc.add_paragraph(f"By {goal['timeframe']}")
        _add_list(doc, "Short Term Objectives:", goal.get("short_term_objectives"))
        _add_list(doc, "Strategies:", goal.get("strategies"))
        _add_list(doc, "Measures / Indicators of Success:", goal.get("measures"))
        _add_list(doc, "Resources:", goal.get("resources"))

    doc.add_heading("Reflection on Success", level=2)
    for q in (
        "How successful have I been in meeting my goals?",
        "How has my professional practice improved?",
        "How has my student learning improved?",
    ):
        doc.add_paragraph(q)
        doc.add_paragraph("")

    doc.add_heading("Review Dates & Comments", level=2)
    for label in ("November:", "January:", "June:"):
        doc.add_paragraph(label)
        doc.add_paragraph("")

    doc.add_paragraph("Teacher signature: ")
    doc.add_paragraph("Principal signature: ")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _add_list(doc, heading, items):
    items = [str(x).strip() for x in (items or []) if str(x).strip()]
    if not items:
        return
    doc.add_paragraph(heading, style="Heading 3")
    for i, item in enumerate(items, start=1):
        doc.add_paragraph(f"{i}. {item}")
